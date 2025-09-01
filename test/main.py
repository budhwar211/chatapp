from __future__ import annotations

import os
import json
import time
import hashlib
import secrets
import logging
import sqlite3
import shutil
from typing import Callable, Dict, List, Optional, Any, Union
from dataclasses import dataclass, field, asdict
from datetime import datetime, timedelta
from pathlib import Path

import requests
from dotenv import load_dotenv

# PDF and DOC generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOC file generation will be unavailable.")

from langgraph.graph import StateGraph, MessagesState, START, END
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import MemorySaver

from langchain_core.tools import tool
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_google_genai import ChatGoogleGenerativeAI
# Removed HuggingFaceEmbeddings due to TensorFlow conflicts

# Vector store + splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# MCP (Model Context Protocol) Integration
try:
    import mcp
    from mcp.client.session import ClientSession
    from mcp.client.stdio import StdioServerParameters, stdio_client
    MCP_AVAILABLE = True
except ImportError:
    MCP_AVAILABLE = False
    print("Warning: MCP not installed. Model Context Protocol features will be unavailable.")


load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# -----------------------------
# Professional Form Generation Classes
# -----------------------------

@dataclass
class FormField:
    """Represents a form field with all necessary properties."""
    name: str
    label: str
    field_type: str
    required: bool = False
    placeholder: str = ""
    options: List[str] = field(default_factory=list)
    validation: str = ""
    description: str = ""
    default_value: str = ""

@dataclass
class FormSection:
    """Represents a section of a form."""
    title: str
    description: str = ""
    fields: List[FormField] = field(default_factory=list)

@dataclass
class ProfessionalForm:
    """Represents a complete professional form."""
    title: str
    description: str
    company_name: str = ""
    form_type: str = "general"
    sections: List[FormSection] = field(default_factory=list)
    footer_text: str = ""
    created_date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    form_id: str = field(default_factory=lambda: secrets.token_hex(4))

class FormGenerator:
    """Professional form generator with PDF and DOC export capabilities."""

    def __init__(self):
        self.output_dir = Path("generated_forms")
        self.output_dir.mkdir(exist_ok=True)

    def create_pdf_form(self, form: ProfessionalForm, filename: str = None) -> str:
        """Generate a professional PDF form."""
        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.pdf"

        filepath = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)

        # Get styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )

        section_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue
        )

        field_style = ParagraphStyle(
            'FieldLabel',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=20
        )

        # Build content
        content = []

        # Header
        if form.company_name:
            company_para = Paragraph(form.company_name, styles['Normal'])
            company_para.alignment = TA_CENTER
            content.append(company_para)
            content.append(Spacer(1, 12))

        # Title
        content.append(Paragraph(form.title, title_style))

        # Description
        if form.description:
            content.append(Paragraph(form.description, styles['Normal']))
            content.append(Spacer(1, 20))

        # Form ID and Date
        info_text = f"Form ID: {form.form_id} | Date: {form.created_date}"
        content.append(Paragraph(info_text, styles['Normal']))
        content.append(Spacer(1, 20))

        # Sections and Fields
        for section in form.sections:
            # Section title
            content.append(Paragraph(section.title, section_style))

            if section.description:
                content.append(Paragraph(section.description, styles['Normal']))
                content.append(Spacer(1, 10))

            # Fields in this section
            for field in section.fields:
                # Field label with required indicator
                label_text = field.label
                if field.required:
                    label_text += " *"

                content.append(Paragraph(label_text, field_style))

                # Field description
                if field.description:
                    desc_style = ParagraphStyle(
                        'FieldDesc',
                        parent=styles['Normal'],
                        fontSize=9,
                        leftIndent=40,
                        textColor=colors.grey
                    )
                    content.append(Paragraph(field.description, desc_style))

                # Input area based on field type
                if field.field_type in ['text', 'email', 'number', 'date']:
                    input_line = "_" * 50
                    content.append(Paragraph(input_line, styles['Normal']))
                elif field.field_type == 'textarea':
                    for _ in range(3):
                        content.append(Paragraph("_" * 70, styles['Normal']))
                elif field.field_type in ['select', 'radio']:
                    for option in field.options:
                        option_text = f"☐ {option}"
                        content.append(Paragraph(option_text, field_style))
                elif field.field_type == 'checkbox':
                    for option in field.options:
                        option_text = f"☐ {option}"
                        content.append(Paragraph(option_text, field_style))

                content.append(Spacer(1, 15))

        # Footer
        if form.footer_text:
            content.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=9,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            content.append(Paragraph(form.footer_text, footer_style))

        # Required fields note
        content.append(Spacer(1, 20))
        required_note = "* Required fields"
        content.append(Paragraph(required_note, styles['Normal']))

        # Build PDF
        doc.build(content)

        logger.info(f"Generated PDF form: {filepath}")
        return str(filepath)

    def create_docx_form(self, form: ProfessionalForm, filename: str = None) -> str:
        """Generate a professional DOCX form."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOC file generation. Install with: pip install python-docx")

        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.docx"

        filepath = self.output_dir / filename

        # Create document
        doc = DocxDocument()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header
        if form.company_name:
            header = doc.add_heading(form.company_name, level=0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_heading(form.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Description
        if form.description:
            desc_para = doc.add_paragraph(form.description)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Form info
        info_para = doc.add_paragraph(f"Form ID: {form.form_id} | Date: {form.created_date}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add spacing
        doc.add_paragraph()

        # Sections and Fields
        for section in form.sections:
            # Section heading
            section_heading = doc.add_heading(section.title, level=2)

            if section.description:
                doc.add_paragraph(section.description)

            # Fields
            for field in section.fields:
                # Field label
                label_text = field.label
                if field.required:
                    label_text += " *"

                field_para = doc.add_paragraph()
                field_run = field_para.add_run(label_text)
                field_run.bold = True

                # Field description
                if field.description:
                    desc_para = doc.add_paragraph(field.description)
                    desc_run = desc_para.runs[0]
                    desc_run.italic = True

                # Input area
                if field.field_type in ['text', 'email', 'number', 'date']:
                    input_para = doc.add_paragraph("_" * 50)
                elif field.field_type == 'textarea':
                    for _ in range(3):
                        doc.add_paragraph("_" * 70)
                elif field.field_type in ['select', 'radio']:
                    for option in field.options:
                        option_para = doc.add_paragraph(f"☐ {option}")
                elif field.field_type == 'checkbox':
                    for option in field.options:
                        option_para = doc.add_paragraph(f"☐ {option}")

                # Add spacing
                doc.add_paragraph()

        # Footer
        if form.footer_text:
            doc.add_page_break()
            footer_para = doc.add_paragraph(form.footer_text)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Required fields note
        required_para = doc.add_paragraph("* Required fields")
        required_run = required_para.runs[0]
        required_run.italic = True

        # Save document
        doc.save(str(filepath))

        logger.info(f"Generated DOCX form: {filepath}")
        return str(filepath)


# -----------------------------
# MCP (Model Context Protocol) Integration
# -----------------------------

@dataclass
class MCPServer:
    """Configuration for an MCP server."""
    name: str
    command: str
    args: List[str] = field(default_factory=list)
    env: Dict[str, str] = field(default_factory=dict)
    enabled: bool = True

class MCPManager:
    """Manages MCP server connections and tool integration."""

    def __init__(self):
        self.servers: Dict[str, MCPServer] = {}
        self.sessions: Dict[str, ClientSession] = {}
        self.mcp_tools: Dict[str, List] = {}

    def register_mcp_server(self, server: MCPServer) -> bool:
        """Register an MCP server configuration."""
        if not MCP_AVAILABLE:
            logger.warning("MCP not available - server registration skipped")
            return False

        self.servers[server.name] = server
        logger.info(f"Registered MCP server: {server.name}")
        return True

    async def connect_server(self, server_name: str) -> bool:
        """Connect to an MCP server and initialize session."""
        if not MCP_AVAILABLE:
            return False

        if server_name not in self.servers:
            logger.error(f"MCP server {server_name} not registered")
            return False

        server = self.servers[server_name]
        if not server.enabled:
            logger.info(f"MCP server {server_name} is disabled")
            return False

        try:
            # Create server parameters
            server_params = StdioServerParameters(
                command=server.command,
                args=server.args,
                env=server.env
            )

            # Connect to server
            async with stdio_client(server_params) as (read, write):
                async with ClientSession(read, write) as session:
                    # Initialize the session
                    await session.initialize()

                    # Get available tools
                    tools_result = await session.list_tools()

                    # Store session and tools
                    self.sessions[server_name] = session
                    self.mcp_tools[server_name] = tools_result.tools if hasattr(tools_result, 'tools') else []

                    logger.info(f"Connected to MCP server {server_name} with {len(self.mcp_tools[server_name])} tools")
                    return True

        except Exception as e:
            logger.error(f"Failed to connect to MCP server {server_name}: {e}")
            return False

    def get_mcp_tools_for_tenant(self, tenant_id: str) -> List:
        """Get MCP tools available for a specific tenant."""
        if not MCP_AVAILABLE:
            return []

        # For now, return all MCP tools - can be enhanced with tenant-specific filtering
        all_tools = []
        for server_name, tools in self.mcp_tools.items():
            all_tools.extend(tools)

        return all_tools

    async def call_mcp_tool(self, server_name: str, tool_name: str, arguments: Dict) -> str:
        """Call an MCP tool on a specific server."""
        if not MCP_AVAILABLE:
            return "MCP not available"

        if server_name not in self.sessions:
            return f"No active session for MCP server {server_name}"

        try:
            session = self.sessions[server_name]
            result = await session.call_tool(tool_name, arguments)

            if hasattr(result, 'content'):
                return str(result.content)
            else:
                return str(result)

        except Exception as e:
            logger.error(f"MCP tool call failed: {e}")
            return f"Error calling MCP tool: {e}"

# Global MCP manager
MCP_MANAGER = MCPManager()

def setup_default_mcp_servers():
    """Set up default MCP servers for common integrations."""
    if not MCP_AVAILABLE:
        return

    # Example MCP servers - can be configured via environment or config file
    default_servers = [
        MCPServer(
            name="filesystem",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-filesystem", "/path/to/allowed/directory"],
            enabled=os.environ.get("MCP_FILESYSTEM_ENABLED", "false").lower() == "true"
        ),
        MCPServer(
            name="git",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-git", "--repository", "."],
            enabled=os.environ.get("MCP_GIT_ENABLED", "false").lower() == "true"
        ),
        MCPServer(
            name="sqlite",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-sqlite", "--db-path", "database.db"],
            enabled=os.environ.get("MCP_SQLITE_ENABLED", "false").lower() == "true"
        )
    ]

    for server in default_servers:
        MCP_MANAGER.register_mcp_server(server)

    logger.info(f"Registered {len(default_servers)} default MCP servers")

# Initialize default MCP servers
setup_default_mcp_servers()


# -----------------------------
# Authentication and Tenant Management
# -----------------------------

@dataclass
class TenantConfig:
    """Configuration for a tenant."""
    tenant_id: str
    name: str
    api_keys: Dict[str, str] = field(default_factory=dict)
    permissions: List[str] = field(default_factory=list)
    rate_limits: Dict[str, float] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    is_active: bool = True
    max_documents: int = 1000
    max_api_calls_per_hour: int = 1000

@dataclass
class DocumentMetadata:
    """Enhanced document metadata for tracking uploaded documents."""
    document_id: str
    filename: str
    file_path: str
    file_size: int
    file_type: str
    upload_timestamp: str
    tenant_id: str
    user_id: Optional[str] = None
    chunk_count: int = 0
    indexed: bool = False
    tags: List[str] = field(default_factory=list)
    file_hash: str = ""
    original_name: str = ""
    description: str = ""

    def __post_init__(self):
        if not self.original_name:
            self.original_name = self.filename
        if not self.file_hash and os.path.exists(self.file_path):
            self.file_hash = self.calculate_file_hash()

    def calculate_file_hash(self) -> str:
        """Calculate SHA256 hash of the file for deduplication."""
        try:
            with open(self.file_path, 'rb') as f:
                return hashlib.sha256(f.read()).hexdigest()
        except Exception:
            return ""

@dataclass
class ChatMessage:
    """Chat message with metadata."""
    message_id: str
    session_id: str
    tenant_id: str
    role: str  # 'user' or 'assistant'
    content: str
    timestamp: str
    user_id: Optional[str] = None
    agent_type: Optional[str] = None
    document_references: List[str] = field(default_factory=list)  # Referenced document IDs

@dataclass
class UserSession:
    """Enhanced user session information."""
    session_id: str
    tenant_id: str
    user_id: Optional[str] = None
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    last_activity: str = field(default_factory=lambda: datetime.now().isoformat())
    permissions: List[str] = field(default_factory=list)
    chat_history: List[ChatMessage] = field(default_factory=list)
    uploaded_documents: List[str] = field(default_factory=list)  # Document IDs

# Enhanced Database Storage System
class DocumentStorage:
    """Enhanced document storage with SQLite backend for persistence."""

    def __init__(self, db_path: str = "document_storage.db"):
        self.db_path = db_path
        self.init_database()

    def init_database(self):
        """Initialize SQLite database with required tables."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        # Documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                document_id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_type TEXT NOT NULL,
                upload_timestamp TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                chunk_count INTEGER DEFAULT 0,
                indexed BOOLEAN DEFAULT FALSE,
                tags TEXT,
                file_hash TEXT,
                original_name TEXT
            )
        ''')

        # Chat messages table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS chat_messages (
                message_id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                timestamp TEXT NOT NULL,
                agent_type TEXT,
                document_references TEXT
            )
        ''')

        # User sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_sessions (
                session_id TEXT PRIMARY KEY,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                created_at TEXT NOT NULL,
                last_activity TEXT NOT NULL,
                permissions TEXT,
                uploaded_documents TEXT
            )
        ''')

        # Create indexes for better performance
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_documents_tenant ON documents(tenant_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_messages_session ON chat_messages(session_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_sessions_tenant ON user_sessions(tenant_id)')

        conn.commit()
        conn.close()

    def save_document(self, doc_metadata: DocumentMetadata) -> bool:
        """Save document metadata to database."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                INSERT OR REPLACE INTO documents
                (document_id, filename, file_path, file_size, file_type, upload_timestamp,
                 tenant_id, user_id, chunk_count, indexed, tags, file_hash, original_name)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc_metadata.document_id,
                doc_metadata.filename,
                doc_metadata.file_path,
                doc_metadata.file_size,
                doc_metadata.file_type,
                doc_metadata.upload_timestamp,
                doc_metadata.tenant_id,
                doc_metadata.user_id,
                doc_metadata.chunk_count,
                doc_metadata.indexed,
                json.dumps(doc_metadata.tags),
                getattr(doc_metadata, 'file_hash', ''),
                getattr(doc_metadata, 'original_name', doc_metadata.filename)
            ))

            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Failed to save document metadata: {e}")
            return False

    def get_documents_by_tenant(self, tenant_id: str) -> List[DocumentMetadata]:
        """Get all documents for a tenant."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                SELECT document_id, filename, file_path, file_size, file_type, upload_timestamp,
                       tenant_id, user_id, chunk_count, indexed, tags, file_hash, original_name
                FROM documents WHERE tenant_id = ?
                ORDER BY upload_timestamp DESC
            ''', (tenant_id,))

            documents = []
            for row in cursor.fetchall():
                doc = DocumentMetadata(
                    document_id=row[0],
                    filename=row[1],
                    file_path=row[2],
                    file_size=row[3],
                    file_type=row[4],
                    upload_timestamp=row[5],
                    tenant_id=row[6],
                    user_id=row[7],
                    chunk_count=row[8] or 0,
                    indexed=bool(row[9]),
                    tags=json.loads(row[10]) if row[10] else []
                )
                # Add additional fields
                doc.file_hash = row[11] or ''
                doc.original_name = row[12] or row[1]
                documents.append(doc)

            conn.close()
            return documents
        except Exception as e:
            logger.error(f"Failed to get documents for tenant {tenant_id}: {e}")
            return []

    def save_chat_message(self, message: ChatMessage) -> bool:
        """Save chat message to database."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                INSERT OR REPLACE INTO chat_messages
                (message_id, session_id, tenant_id, user_id, role, content, timestamp, agent_type, document_references)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                message.message_id,
                message.session_id,
                message.tenant_id,
                message.user_id,
                message.role,
                message.content,
                message.timestamp,
                message.agent_type,
                json.dumps(message.document_references)
            ))

            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Failed to save chat message: {e}")
            return False

    def get_chat_history(self, session_id: str, limit: int = 50) -> List[ChatMessage]:
        """Get chat history for a session."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                SELECT message_id, session_id, tenant_id, user_id, role, content, timestamp, agent_type, document_references
                FROM chat_messages WHERE session_id = ?
                ORDER BY timestamp DESC LIMIT ?
            ''', (session_id, limit))

            messages = []
            for row in cursor.fetchall():
                message = ChatMessage(
                    message_id=row[0],
                    session_id=row[1],
                    tenant_id=row[2],
                    role=row[4],
                    content=row[5],
                    timestamp=row[6],
                    user_id=row[3],
                    agent_type=row[7],
                    document_references=json.loads(row[8]) if row[8] else []
                )
                messages.append(message)

            conn.close()
            return list(reversed(messages))  # Return in chronological order
        except Exception as e:
            logger.error(f"Failed to get chat history for session {session_id}: {e}")
            return []

# Global storage instance
document_storage = DocumentStorage()

# Global runtime context
CURRENT_TENANT_ID: Optional[str] = None
CURRENT_SESSION: Optional[UserSession] = None

def set_current_tenant(tenant_id: str):
    """Set the current tenant for the session"""
    global CURRENT_TENANT_ID, CURRENT_SESSION
    CURRENT_TENANT_ID = tenant_id
    if tenant_id in _tenant_registry:
        CURRENT_SESSION = create_session(tenant_id)
    logger.info(f"Set current tenant to: {tenant_id}")

# Tenant registry
_tenant_registry: Dict[str, TenantConfig] = {}
_active_sessions: Dict[str, UserSession] = {}

def create_tenant(tenant_id: str, name: str, permissions: Optional[List[str]] = None) -> TenantConfig:
    """Create a new tenant with default configuration."""
    if tenant_id in _tenant_registry:
        raise ValueError(f"Tenant {tenant_id} already exists")
    
    config = TenantConfig(
        tenant_id=tenant_id,
        name=name,
        permissions=permissions or ["read_documents", "use_tools", "generate_forms"],
        rate_limits={"default": 0.5, "search_web": 1.0, "get_weather": 0.5}
    )
    
    _tenant_registry[tenant_id] = config
    logger.info(f"Created tenant: {tenant_id}")
    return config

def get_tenant_config(tenant_id: str) -> Optional[TenantConfig]:
    """Get tenant configuration."""
    return _tenant_registry.get(tenant_id)

def authenticate_tenant(tenant_id: str, api_key: Optional[str] = None) -> bool:
    """Authenticate a tenant (simplified for demo)."""
    config = get_tenant_config(tenant_id)
    if not config or not config.is_active:
        return False
    
    # In production, you'd validate the API key here
    if api_key and "master_key" in config.api_keys:
        return config.api_keys["master_key"] == api_key
    
    # For demo purposes, allow access without API key
    return True

def create_session(tenant_id: str, user_id: Optional[str] = None) -> UserSession:
    """Create a new user session."""
    if not authenticate_tenant(tenant_id):
        raise ValueError(f"Invalid tenant: {tenant_id}")
    
    session_id = secrets.token_urlsafe(32)
    config = get_tenant_config(tenant_id)
    
    session = UserSession(
        session_id=session_id,
        tenant_id=tenant_id,
        user_id=user_id,
        permissions=config.permissions if config else []
    )
    
    _active_sessions[session_id] = session
    logger.info(f"Created session {session_id} for tenant {tenant_id}")
    return session

def get_session(session_id: str) -> Optional[UserSession]:
    """Get session by ID."""
    session = _active_sessions.get(session_id)
    if session:
        # Update last activity
        session.last_activity = datetime.now().isoformat()
    return session

def has_permission(permission: str, session: Optional[UserSession] = None) -> bool:
    """Check if current session has permission."""
    current_session = session or CURRENT_SESSION
    if not current_session:
        return False
    return permission in current_session.permissions

def require_permission(permission: str) -> Callable:
    """Decorator to require specific permission."""
    def decorator(func: Callable) -> Callable:
        def wrapper(*args, **kwargs):
            if not has_permission(permission):
                raise PermissionError(f"Permission required: {permission}")
            return func(*args, **kwargs)
        return wrapper
    return decorator

# Initialize default tenant
def initialize_default_tenant():
    """Initialize default tenant for demo purposes."""
    if "default" not in _tenant_registry:
        create_tenant(
            "default",
            "Default Tenant",
            ["read_documents", "use_tools", "generate_forms", "admin"]
        )

initialize_default_tenant()


# -----------------------------
# Enhanced Dynamic Tooling Infrastructure
# -----------------------------

from collections import defaultdict
from threading import Lock
import logging

# Enhanced tool registry with metadata
_dynamic_tool_registry: Dict[str, List] = {}
_tool_metadata: Dict[str, Dict] = {}
_last_call_timestamp_per_tool: Dict[str, float] = {}
_tool_call_counts: Dict[str, int] = defaultdict(int)
_tool_error_counts: Dict[str, int] = defaultdict(int)
_registry_lock = Lock()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _rate_limited(tool_name: str, min_interval_seconds: float = 0.5) -> bool:
    """Enhanced rate limiting with per-tool configuration."""
    now = time.time()
    last = _last_call_timestamp_per_tool.get(tool_name)
    
    # Get tool-specific rate limit if available
    tool_meta = _tool_metadata.get(tool_name, {})
    interval = tool_meta.get('rate_limit_seconds', min_interval_seconds)
    
    if last is None or now - last >= interval:
        _last_call_timestamp_per_tool[tool_name] = now
        return True
    return False


def get_tenant_tools(tenant_id: Optional[str]) -> List:
    """Get all available tools for a tenant with enhanced filtering and MCP integration."""
    base_tools = [search_web, get_weather, get_document_stats_tool]
    tenant_list = _dynamic_tool_registry.get(tenant_id or "default", [])

    # Filter out disabled tools - handle both function tools and StructuredTool objects
    active_tenant_tools = []
    for tool in tenant_list:
        tool_name = getattr(tool, 'name', getattr(tool, '__name__', str(tool)))
        is_enabled = _tool_metadata.get(tool_name, {}).get('enabled', True)
        if is_enabled:
            active_tenant_tools.append(tool)

    # Add MCP tools if available
    mcp_tools = MCP_MANAGER.get_mcp_tools_for_tenant(tenant_id or "default")

    return base_tools + active_tenant_tools + mcp_tools


def register_dynamic_tool(tenant_id: str, dynamic_tool, metadata: Optional[Dict] = None) -> None:
    """Enhanced tool registration with metadata and validation."""
    with _registry_lock:
        tools_for_tenant = _dynamic_tool_registry.setdefault(tenant_id, [])
        
        # Get tool name safely for both function tools and StructuredTool objects
        tool_name = getattr(dynamic_tool, 'name', getattr(dynamic_tool, '__name__', str(dynamic_tool)))
        
        # Replace by name if already exists
        existing_names = set()
        for tool in tools_for_tenant:
            existing_tool_name = getattr(tool, 'name', getattr(tool, '__name__', str(tool)))
            existing_names.add(existing_tool_name)
            
        if tool_name in existing_names:
            # Remove existing tool with the same name
            tools_for_tenant[:] = [
                t for t in tools_for_tenant 
                if getattr(t, 'name', getattr(t, '__name__', str(t))) != tool_name
            ]
        
        tools_for_tenant.append(dynamic_tool)
        
        # Store metadata
        tool_meta = metadata or {}
        tool_meta.update({
            'tenant_id': tenant_id,
            'registered_at': datetime.now().isoformat(),
            'enabled': tool_meta.get('enabled', True),
            'rate_limit_seconds': tool_meta.get('rate_limit_seconds', 0.5),
            'max_retries': tool_meta.get('max_retries', 3)
        })
        _tool_metadata[tool_name] = tool_meta
        
        logger.info(f"Registered tool '{tool_name}' for tenant '{tenant_id}'")


def unregister_tool(tenant_id: str, tool_name: str) -> bool:
    """Remove a tool from a tenant's registry."""
    with _registry_lock:
        tools_for_tenant = _dynamic_tool_registry.get(tenant_id, [])
        original_count = len(tools_for_tenant)
        
        # Handle both function tools and StructuredTool objects
        _dynamic_tool_registry[tenant_id] = [
            t for t in tools_for_tenant 
            if getattr(t, 'name', getattr(t, '__name__', str(t))) != tool_name
        ]
        
        if tool_name in _tool_metadata:
            del _tool_metadata[tool_name]
        
        removed = len(_dynamic_tool_registry[tenant_id]) < original_count
        if removed:
            logger.info(f"Unregistered tool '{tool_name}' from tenant '{tenant_id}'")
        
        return removed


def get_tool_stats(tenant_id: Optional[str] = None) -> Dict:
    """Get statistics about tool usage."""
    if tenant_id:
        tools = get_tenant_tools(tenant_id)
        # Handle both function tools and StructuredTool objects
        tool_names = []
        for tool in tools:
            tool_name = getattr(tool, 'name', getattr(tool, '__name__', str(tool)))
            tool_names.append(tool_name)
    else:
        tool_names = list(_tool_metadata.keys())
    
    stats = {}
    for tool_name in tool_names:
        stats[tool_name] = {
            'call_count': _tool_call_counts.get(tool_name, 0),
            'error_count': _tool_error_counts.get(tool_name, 0),
            'last_called': _last_call_timestamp_per_tool.get(tool_name),
            'metadata': _tool_metadata.get(tool_name, {})
        }
    
    return stats


def make_http_get_tool(
    name: str,
    description: str,
    base_url_env: str,
    api_key_env: Optional[str] = None,
    rate_limit_seconds: float = 0.5,
    timeout: int = 20
) -> any:
    """Enhanced HTTP GET tool with better error handling and configuration."""

    from langchain_core.tools import StructuredTool
    from pydantic import BaseModel, Field

    class InputSchema(BaseModel):
        path: str = Field(..., description="URL path to append to the base URL, starting with '/'")
        query: Optional[Dict[str, str]] = Field(default=None, description="Query params as key-value map")
        headers: Optional[Dict[str, str]] = Field(default=None, description="Additional headers")

    base_url = os.environ.get(base_url_env)
    api_key = os.environ.get(api_key_env) if api_key_env else None

    def _run(path: str, query: Optional[Dict[str, str]] = None, headers: Optional[Dict[str, str]] = None) -> str:
        tool_name = f"{name}"
        
        # Track call count
        _tool_call_counts[tool_name] += 1
        
        # Rate limiting
        if not _rate_limited(tool_name, rate_limit_seconds):
            return "Rate limited. Please retry shortly."
        
        # Validation
        if not base_url:
            _tool_error_counts[tool_name] += 1
            return f"HTTP GET tool misconfigured: missing env {base_url_env}"
        
        # Build request
        url = base_url.rstrip("/") + path
        request_headers = {"Accept": "application/json", "User-Agent": "Multi-Agent-Chatbot/1.0"}
        
        if api_key:
            request_headers["Authorization"] = f"Bearer {api_key}"
        
        if headers:
            request_headers.update(headers)
        
        try:
            logger.info(f"Making HTTP GET request to {url}")
            resp = requests.get(
                url,
                params=query or {},
                headers=request_headers,
                timeout=timeout
            )
            
            if resp.ok:
                content = resp.text[:4000]  # Limit response size
                logger.info(f"HTTP GET successful for {tool_name}")
                return content
            else:
                error_msg = f"HTTP {resp.status_code}: {resp.text[:800]}"
                logger.warning(f"HTTP GET failed for {tool_name}: {error_msg}")
                _tool_error_counts[tool_name] += 1
                return error_msg
                
        except requests.exceptions.Timeout:
            error_msg = f"HTTP GET timeout after {timeout}s"
            logger.error(f"HTTP GET timeout for {tool_name}")
            _tool_error_counts[tool_name] += 1
            return error_msg
        except Exception as exc:
            error_msg = f"HTTP GET error: {exc}"
            logger.error(f"HTTP GET error for {tool_name}: {exc}")
            _tool_error_counts[tool_name] += 1
            return error_msg

    return StructuredTool.from_function(
        name=name,
        description=description,
        func=_run,
        args_schema=InputSchema,
    )


def make_http_post_tool(
    name: str,
    description: str,
    base_url_env: str,
    api_key_env: Optional[str] = None,
    rate_limit_seconds: float = 1.0,
    timeout: int = 30
) -> any:
    """Create an HTTP POST tool for API interactions."""
    
    from langchain_core.tools import StructuredTool
    from pydantic import BaseModel, Field

    class InputSchema(BaseModel):
        path: str = Field(..., description="URL path to append to the base URL")
        data: Optional[Dict] = Field(default=None, description="JSON data to send in request body")
        headers: Optional[Dict[str, str]] = Field(default=None, description="Additional headers")

    base_url = os.environ.get(base_url_env)
    api_key = os.environ.get(api_key_env) if api_key_env else None

    def _run(path: str, data: Optional[Dict] = None, headers: Optional[Dict[str, str]] = None) -> str:
        tool_name = f"{name}"
        
        _tool_call_counts[tool_name] += 1
        
        if not _rate_limited(tool_name, rate_limit_seconds):
            return "Rate limited. Please retry shortly."
        
        if not base_url:
            _tool_error_counts[tool_name] += 1
            return f"HTTP POST tool misconfigured: missing env {base_url_env}"
        
        url = base_url.rstrip("/") + path
        request_headers = {"Content-Type": "application/json", "User-Agent": "Multi-Agent-Chatbot/1.0"}
        
        if api_key:
            request_headers["Authorization"] = f"Bearer {api_key}"
        
        if headers:
            request_headers.update(headers)
        
        try:
            logger.info(f"Making HTTP POST request to {url}")
            resp = requests.post(
                url,
                json=data,
                headers=request_headers,
                timeout=timeout
            )
            
            if resp.ok:
                content = resp.text[:4000]
                logger.info(f"HTTP POST successful for {tool_name}")
                return content
            else:
                error_msg = f"HTTP {resp.status_code}: {resp.text[:800]}"
                logger.warning(f"HTTP POST failed for {tool_name}: {error_msg}")
                _tool_error_counts[tool_name] += 1
                return error_msg
                
        except Exception as exc:
            error_msg = f"HTTP POST error: {exc}"
            logger.error(f"HTTP POST error for {tool_name}: {exc}")
            _tool_error_counts[tool_name] += 1
            return error_msg

    return StructuredTool.from_function(
        name=name,
        description=description,
        func=_run,
        args_schema=InputSchema,
    )


# -----------------------------
# Enhanced Built-in Tools
# -----------------------------


@tool
def search_web(query: str) -> str:
    """Enhanced web search via DuckDuckGo instant answers; returns brief summary/snippets."""
    tool_name = "search_web"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 1.0):  # 1 second rate limit for web search
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info(f"Performing web search for: {query}")
        resp = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1},
            timeout=12,
        )
        data = resp.json() if resp.ok else {}
        
        # Try abstract first
        abstract = data.get("AbstractText") or data.get("Abstract") or ""
        if abstract:
            logger.info(f"Web search successful for: {query}")
            return f"Abstract: {abstract}"
        
        # Try related topics
        related = data.get("RelatedTopics", [])
        snippets: List[str] = []
        for item in related:
            if isinstance(item, dict) and item.get("Text"):
                snippets.append(item["Text"])
            elif isinstance(item, dict) and item.get("Topics"):
                for sub in item.get("Topics", []):
                    if sub.get("Text"):
                        snippets.append(sub["Text"])
            if len(snippets) >= 3:
                break
        
        if snippets:
            result = "Related info: " + " | ".join(snippets[:3])
            logger.info(f"Web search successful for: {query}")
            return result
        
        # Try definition
        definition = data.get("Definition", "")
        if definition:
            return f"Definition: {definition}"
        
        logger.warning(f"No results found for web search: {query}")
        return "No quick answer found. Try rephrasing your query."
        
    except Exception as exc:
        logger.error(f"Web search failed for {query}: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Search failed: {exc}"


@tool
def get_weather(city: str) -> str:
    """Enhanced weather lookup for a city using Open‑Meteo API."""
    tool_name = "get_weather"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 0.5):
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info(f"Getting weather for: {city}")
        
        # Geocoding
        geo = requests.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": city, "count": 1, "language": "en", "format": "json"},
            timeout=12,
        )
        geo_data = geo.json() if geo.ok else {}
        results = geo_data.get("results") or []
        
        if not results:
            logger.warning(f"City not found: {city}")
            return f"Could not find city '{city}'. Please check the spelling."
        
        loc = results[0]
        lat, lon = loc["latitude"], loc["longitude"]
        place = f"{loc.get('name')}, {loc.get('country_code', '')}".strip()

        # Weather data
        w = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": lat,
                "longitude": lon,
                "current": "temperature_2m,precipitation,relative_humidity_2m,apparent_temperature,is_day,weather_code,wind_speed_10m",
                "timezone": "auto"
            },
            timeout=12,
        )
        w_data = w.json() if w.ok else {}
        cur = w_data.get("current") or {}
        
        temp = cur.get("temperature_2m")
        precip = cur.get("precipitation", 0)
        humidity = cur.get("relative_humidity_2m")
        feels = cur.get("apparent_temperature")
        wind_speed = cur.get("wind_speed_10m", 0)
        is_day = cur.get("is_day", 1)
        
        time_of_day = "day" if is_day else "night"
        
        result = (
            f"Weather in {place} ({time_of_day}): "
            f"temp {temp}°C (feels like {feels}°C), "
            f"humidity {humidity}%, wind {wind_speed} km/h"
        )
        
        if precip > 0:
            result += f", precipitation {precip} mm"
        
        logger.info(f"Weather lookup successful for: {city}")
        return result
        
    except Exception as exc:
        logger.error(f"Weather lookup failed for {city}: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Weather lookup failed: {exc}"


@tool
def get_document_stats_tool(tenant_id: Optional[str] = None) -> str:
    """Get statistics about indexed documents for the current or specified tenant."""
    current_tenant = tenant_id or CURRENT_TENANT_ID or "default"
    
    tool_name = "get_document_stats"
    _tool_call_counts[tool_name] += 1
    
    try:
        stats = get_document_stats(current_tenant)
        
        if "error" in stats:
            return f"Error getting document stats: {stats['error']}"
        
        result = f"Document Statistics for tenant '{current_tenant}':\n"
        result += f"- Total chunks: {stats['total_chunks']}\n"
        result += f"- Unique sources: {stats['unique_sources']}\n"
        result += f"- File types: {', '.join([f'{k}({v})' for k, v in stats['file_types'].items()])}\n"
        
        if stats['sample_sources']:
            result += f"- Sample sources: {', '.join(stats['sample_sources'][:3])}..."
        
        return result
        
    except Exception as exc:
        logger.error(f"Error getting document stats: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Error getting document stats: {exc}"


@tool
def get_tool_statistics() -> str:
    """Get usage statistics for all tools."""
    tool_name = "get_tool_statistics"
    _tool_call_counts[tool_name] += 1
    
    try:
        stats = get_tool_stats()
        
        if not stats:
            return "No tool usage statistics available."
        
        result = "Tool Usage Statistics:\n"
        for name, data in stats.items():
            calls = data['call_count']
            errors = data['error_count']
            success_rate = ((calls - errors) / calls * 100) if calls > 0 else 0
            
            result += f"- {name}: {calls} calls, {errors} errors ({success_rate:.1f}% success)\n"
        
        return result
        
    except Exception as exc:
        logger.error(f"Error getting tool statistics: {exc}")
        return f"Error getting tool statistics: {exc}"


# -----------------------------
# LLM setup
# -----------------------------


def get_llm(temperature: float = 0) -> ChatGoogleGenerativeAI:
    model = os.environ.get("GOOGLE_MODEL", "gemini-2.0-flash")
    return ChatGoogleGenerativeAI(model=model, temperature=temperature)


def build_llm_with_tools_for_tenant(tenant_id: Optional[str]):
    tools = get_tenant_tools(tenant_id)
    return get_llm(temperature=0).bind_tools(tools)


# -----------------------------
# Enhanced RAG: Indexing and Retrieval per Tenant (FAISS)
# -----------------------------

import hashlib
from datetime import datetime
from pathlib import Path

# Simple embeddings implementation without TensorFlow dependencies
class SimpleEmbeddings(Embeddings):
    """Simple embeddings using basic text processing without heavy dependencies"""

    def __init__(self):
        self.dimension = 384

    def _text_to_vector(self, text):
        """Convert text to a simple vector representation"""
        import hashlib
        import re

        # Clean and normalize text
        text = re.sub(r'[^\w\s]', '', text.lower())
        words = text.split()

        # Create features based on text characteristics
        features = []

        # 1. Hash-based features (first 100 dimensions)
        hash_obj = hashlib.md5(text.encode())
        hash_hex = hash_obj.hexdigest()
        hash_features = [int(hash_hex[i:i+2], 16) / 255.0 for i in range(0, min(len(hash_hex), 50), 2)]
        features.extend(hash_features)

        # 2. Length-based features
        features.extend([
            len(text) / 1000.0,  # Text length
            len(words) / 100.0,  # Word count
            sum(len(w) for w in words) / max(len(words), 1) / 10.0,  # Average word length
        ])

        # 3. Character frequency features (26 letters)
        char_freq = [0] * 26
        for char in text:
            if 'a' <= char <= 'z':
                char_freq[ord(char) - ord('a')] += 1
        total_chars = sum(char_freq)
        if total_chars > 0:
            char_freq = [f / total_chars for f in char_freq]
        features.extend(char_freq)

        # 4. Word pattern features
        if words:
            features.extend([
                sum(1 for w in words if len(w) > 5) / len(words),  # Long words ratio
                sum(1 for w in words if w.isupper()) / len(words),  # Uppercase words ratio
                sum(1 for w in words if w.isdigit()) / len(words),  # Numeric words ratio
            ])
        else:
            features.extend([0, 0, 0])

        # 5. Pad or truncate to desired dimension
        while len(features) < self.dimension:
            features.extend(features[:min(self.dimension - len(features), len(features))])

        return features[:self.dimension]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        """Embed a list of documents"""
        return [self._text_to_vector(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        """Embed a single query"""
        return self._text_to_vector(text)

# Initialize embeddings
EMBEDDINGS = SimpleEmbeddings()
logger.info("Using simple embeddings implementation (no TensorFlow dependencies)")

# Global form generator instance
FORM_GENERATOR = FormGenerator()

def _json_to_professional_form(form_data: dict) -> ProfessionalForm:
    """Convert JSON form data to ProfessionalForm object."""
    # Extract basic info
    title = form_data.get("title", "Untitled Form")
    description = form_data.get("description", "")
    company_name = form_data.get("company_name", "")
    form_type = form_data.get("form_type", "general")
    footer_text = form_data.get("footer_text", "")

    # Handle both old format (direct fields) and new format (sections)
    sections = []

    if "sections" in form_data and isinstance(form_data["sections"], list):
        # New format with sections
        for section_data in form_data["sections"]:
            section_title = section_data.get("title", "Section")
            section_desc = section_data.get("description", "")

            fields = []
            for field_data in section_data.get("fields", []):
                field = FormField(
                    name=field_data.get("name", ""),
                    label=field_data.get("label", ""),
                    field_type=field_data.get("field_type", field_data.get("type", "text")),
                    required=field_data.get("required", False),
                    placeholder=field_data.get("placeholder", ""),
                    options=field_data.get("options", []),
                    validation=field_data.get("validation", ""),
                    description=field_data.get("description", ""),
                    default_value=field_data.get("default_value", "")
                )
                fields.append(field)

            section = FormSection(
                title=section_title,
                description=section_desc,
                fields=fields
            )
            sections.append(section)

    elif "fields" in form_data and isinstance(form_data["fields"], list):
        # Old format with direct fields - create a single section
        fields = []
        for field_data in form_data["fields"]:
            field = FormField(
                name=field_data.get("name", ""),
                label=field_data.get("label", ""),
                field_type=field_data.get("field_type", field_data.get("type", "text")),
                required=field_data.get("required", False),
                placeholder=field_data.get("placeholder", ""),
                options=field_data.get("options", []),
                validation=field_data.get("validation", ""),
                description=field_data.get("description", ""),
                default_value=field_data.get("default_value", "")
            )
            fields.append(field)

        # Create a single section
        section = FormSection(
            title="Form Fields",
            description="",
            fields=fields
        )
        sections.append(section)

    # Create the professional form
    professional_form = ProfessionalForm(
        title=title,
        description=description,
        company_name=company_name,
        form_type=form_type,
        sections=sections,
        footer_text=footer_text
    )

    return professional_form


def _tenant_index_path(tenant_id: str) -> str:
    os.makedirs("indices", exist_ok=True)
    return os.path.join("indices", f"faiss_{tenant_id}")


def _get_file_hash(file_path: str) -> str:
    """Generate hash for file content to detect changes."""
    try:
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except Exception:
        return ""


def _extract_text_from_file(file_path: str) -> tuple[str, dict]:
    """Enhanced text extraction with better metadata."""
    path_obj = Path(file_path)
    ext = path_obj.suffix.lower()
    
    metadata = {
        "source": file_path,
        "filename": path_obj.name,
        "file_type": ext,
        "file_size": path_obj.stat().st_size if path_obj.exists() else 0,
        "modified_time": datetime.fromtimestamp(path_obj.stat().st_mtime).isoformat() if path_obj.exists() else "",
        "file_hash": _get_file_hash(file_path)
    }
    
    text = ""
    
    try:
        if ext in {".txt", ".md", ".csv"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(file_path)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                metadata["page_count"] = len(reader.pages)
            except Exception as exc:  # noqa: BLE001
                text = f"[PDF read error: {exc}]"
                metadata["error"] = str(exc)
        elif ext == ".docx":
            try:
                import docx  # type: ignore
                d = docx.Document(file_path)
                text = "\n".join(p.text for p in d.paragraphs)
                metadata["paragraph_count"] = len(d.paragraphs)
            except Exception as exc:  # noqa: BLE001
                text = f"[DOCX read error: {exc}]"
                metadata["error"] = str(exc)
        elif ext == ".json":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    text = json.dumps(data, indent=2)
                metadata["json_keys"] = list(data.keys()) if isinstance(data, dict) else []
            except Exception as exc:
                text = f"[JSON read error: {exc}]"
                metadata["error"] = str(exc)
        else:
            # Try to read as text for other extensions
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception:
                text = f"[Unsupported file type: {ext}]"
                
    except Exception as exc:
        text = f"[File read error: {exc}]"
        metadata["error"] = str(exc)
    
    return text, metadata


def ingest_single_document(tenant_id: str, file_path: str, user_id: Optional[str] = None,
                          chunk_size: int = 1000, chunk_overlap: int = 150) -> Dict[str, Any]:
    """Enhanced single document ingestion with metadata tracking."""
    try:
        # Check if file already exists (deduplication)
        file_hash = hashlib.sha256(open(file_path, 'rb').read()).hexdigest()
        existing_docs = document_storage.get_documents_by_tenant(tenant_id)

        for doc in existing_docs:
            if doc.file_hash == file_hash:
                return {
                    "success": True,
                    "message": f"Document already exists: {doc.filename}",
                    "document_id": doc.document_id,
                    "duplicate": True
                }

        # Extract text and metadata
        text, base_metadata = _extract_text_from_file(file_path)

        if not text.strip():
            return {"success": False, "message": "No text content found in document"}

        # Create document metadata
        document_id = secrets.token_urlsafe(16)
        file_stat = os.stat(file_path)

        doc_metadata = DocumentMetadata(
            document_id=document_id,
            filename=os.path.basename(file_path),
            file_path=file_path,
            file_size=file_stat.st_size,
            file_type=Path(file_path).suffix.lower(),
            upload_timestamp=datetime.now().isoformat(),
            tenant_id=tenant_id,
            user_id=user_id,
            file_hash=file_hash,
            original_name=os.path.basename(file_path)
        )

        # Split text into chunks
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        chunks = splitter.split_text(text)
        doc_metadata.chunk_count = len(chunks)

        # Create documents for vector store
        docs: List[Document] = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "tenant_id": tenant_id,
                "document_id": document_id,
                "chunk_id": i,
                "chunk_count": len(chunks),
                "chunk_size": len(chunk),
                "ingestion_time": datetime.now().isoformat()
            })

            # Sanitize metadata to ensure all values are serializable
            sanitized_metadata = {}
            for key, value in chunk_metadata.items():
                try:
                    # Convert to string if not a basic type
                    if isinstance(value, (str, int, float, bool, type(None))):
                        sanitized_metadata[key] = value
                    else:
                        sanitized_metadata[key] = str(value)
                except Exception:
                    sanitized_metadata[key] = "unknown"

            try:
                # Create document with error handling for Pydantic compatibility
                doc = Document(
                    page_content=chunk,
                    metadata=sanitized_metadata
                )
                docs.append(doc)
            except Exception as doc_error:
                logger.error(f"Error creating document for chunk {i}: {doc_error}")
                # Try with minimal metadata as fallback
                try:
                    doc = Document(
                        page_content=chunk,
                        metadata={"source": sanitized_metadata.get("source", "unknown")}
                    )
                    docs.append(doc)
                except Exception as fallback_error:
                    logger.error(f"Fallback document creation also failed: {fallback_error}")
                    continue

        # Save to vector store
        index_dir = _tenant_index_path(tenant_id)
        try:
            logger.info(f"Attempting to save {len(docs)} documents to vector store at {index_dir}")

            # Debug: Check document structure
            if docs:
                sample_doc = docs[0]
                logger.info(f"Sample document type: {type(sample_doc)}")
                logger.info(f"Sample metadata keys: {list(sample_doc.metadata.keys())}")

            if os.path.isdir(index_dir):
                logger.info("Loading existing vector store")
                try:
                    vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
                    logger.info("Adding documents to existing vector store")
                    vs.add_documents(docs)
                except (KeyError, AttributeError, Exception) as load_error:
                    logger.warning(f"Failed to load existing vector store (likely version incompatibility): {load_error}")
                    logger.info("Creating new vector store to replace corrupted one")
                    # Remove corrupted index directory
                    import shutil
                    shutil.rmtree(index_dir, ignore_errors=True)
                    vs = FAISS.from_documents(docs, EMBEDDINGS)
            else:
                logger.info("Creating new vector store")
                vs = FAISS.from_documents(docs, EMBEDDINGS)

            logger.info("Saving vector store to disk")
            vs.save_local(index_dir)
            logger.info("Vector store saved successfully")

            doc_metadata.indexed = True
        except Exception as e:
            logger.error(f"Failed to save to vector store: {e}")
            logger.error(f"Error type: {type(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return {"success": False, "message": f"Vector indexing failed: {e}"}

        # Save document metadata to database
        if document_storage.save_document(doc_metadata):
            return {
                "success": True,
                "message": f"Document processed successfully: {doc_metadata.filename}",
                "document_id": document_id,
                "chunks": len(chunks),
                "duplicate": False
            }
        else:
            return {"success": False, "message": "Failed to save document metadata"}

    except Exception as e:
        logger.error(f"Error processing document {file_path}: {e}")
        return {"success": False, "message": f"Processing failed: {e}"}

def ingest_multiple_documents(tenant_id: str, file_paths: List[str], user_id: Optional[str] = None) -> Dict[str, Any]:
    """Process multiple documents simultaneously."""
    results = []
    successful = 0
    failed = 0
    duplicates = 0

    for file_path in file_paths:
        result = ingest_single_document(tenant_id, file_path, user_id)
        results.append({
            "file_path": file_path,
            "filename": os.path.basename(file_path),
            **result
        })

        if result["success"]:
            if result.get("duplicate", False):
                duplicates += 1
            else:
                successful += 1
        else:
            failed += 1

    return {
        "success": True,
        "total_files": len(file_paths),
        "successful": successful,
        "failed": failed,
        "duplicates": duplicates,
        "results": results
    }

def ingest_documents_from_dir(tenant_id: str, source_dir: str, chunk_size: int = 1000, chunk_overlap: int = 150) -> str:
    """Enhanced document ingestion with better processing and metadata."""
    file_paths = []

    for root, _dirs, files in os.walk(source_dir):
        for fname in files:
            file_path = os.path.join(root, fname)

            # Skip hidden files and common non-document files
            if fname.startswith('.') or fname.lower().endswith(('.exe', '.dll', '.so', '.dylib')):
                continue

            file_paths.append(file_path)

    if not file_paths:
        return "No documents found to ingest."

    result = ingest_multiple_documents(tenant_id, file_paths)
    return f"Processed {result['total_files']} files: {result['successful']} successful, {result['failed']} failed, {result['duplicates']} duplicates"


def get_retriever_for_tenant(tenant_id: str):
    """Enhanced retriever with better search capabilities."""
    index_dir = _tenant_index_path(tenant_id)
    if not os.path.isdir(index_dir):
        return None
        
    try:
        vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
    except (KeyError, AttributeError, Exception) as exc:
        logger.warning(f"Vector store for tenant {tenant_id} is corrupted (likely version incompatibility): {exc}")
        logger.info(f"Removing corrupted vector store at {index_dir}")
        import shutil
        shutil.rmtree(index_dir, ignore_errors=True)
        return None
    except Exception as exc:
        logger.error(f"Error loading vector store for tenant {tenant_id}: {exc}")
        return None
    
    def _retrieve(query: str, k: int = 4, score_threshold: float = 0.7) -> List[Document]:
        """Enhanced retrieval with similarity scoring."""
        try:
            # Use similarity search with score
            docs_with_scores = vs.similarity_search_with_score(query, k=k*2)  # Get more initially
            
            # Filter by score threshold and limit results
            filtered_docs = [
                doc for doc, score in docs_with_scores
                if score <= score_threshold  # Lower score = higher similarity in FAISS
            ][:k]
            
            # If no docs meet threshold, return top k anyway
            if not filtered_docs:
                filtered_docs = [doc for doc, _ in docs_with_scores[:k]]
            
            return filtered_docs
            
        except Exception as exc:
            print(f"Error during retrieval: {exc}")
            return []

    return _retrieve


def get_document_stats(tenant_id: str) -> dict:
    """Get statistics about indexed documents for a tenant."""
    index_dir = _tenant_index_path(tenant_id)
    if not os.path.isdir(index_dir):
        return {"error": "No index found for tenant"}
    
    try:
        vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
        
        # Get basic stats
        total_chunks = vs.index.ntotal
        
        # Sample some documents to get metadata stats
        sample_docs = vs.similarity_search("", k=min(100, total_chunks)) if total_chunks > 0 else []
        
        file_types = {}
        sources = set()
        
        for doc in sample_docs:
            metadata = doc.metadata
            file_type = metadata.get("file_type", "unknown")
            file_types[file_type] = file_types.get(file_type, 0) + 1
            sources.add(metadata.get("source", "unknown"))
        
        return {
            "tenant_id": tenant_id,
            "total_chunks": total_chunks,
            "unique_sources": len(sources),
            "file_types": file_types,
            "sample_sources": list(sources)[:10]  # Show first 10 sources
        }
        
    except Exception as exc:
        return {"error": f"Error getting stats: {exc}"}


# -----------------------------
# Agent Nodes
# -----------------------------


def node_router(state: MessagesState) -> str:
    """Enhanced router with better intent classification."""
    last_user = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            last_user = getattr(msg, "content", "")
            break

    # Enhanced routing logic with more context
    prompt = (
        "You are an intelligent router for a multi-agent chatbot system. "
        "Analyze the user's message and classify their intent into one of these categories:\n\n"
        "- greeting: greetings, small talk, general conversation, introductions\n"
        "- doc_qa: questions about documents, files, or knowledge base content\n"
        "- api_exec: requests to perform actions, call APIs, get external data (weather, search, etc.)\n"
        "- form_gen: requests to create forms, collect structured data, or generate input fields\n"
        "- analytics: requests for data analysis, statistics, insights, reports, or metrics\n"
        "- escalate: requests for human help, complaints, or complex issues beyond AI capability\n\n"
        "Consider context clues like:\n"
        "- Keywords related to documents, files, or knowledge\n"
        "- Action words like 'get', 'fetch', 'call', 'search'\n"
        "- Form-related terms like 'form', 'input', 'collect', 'survey'\n"
        "- Analytics terms like 'analyze', 'statistics', 'metrics', 'report', 'insights'\n"
        "- Escalation phrases like 'human', 'agent', 'help', 'support'\n\n"
        "Respond with only the category name."
    )
    
    llm = get_llm(temperature=0)
    res = llm.invoke([("system", prompt), ("user", last_user or "hello")])
    label = (getattr(res, "content", "") or "").strip().lower()
    
    # Fallback logic with keyword detection
    if label not in {"greeting", "doc_qa", "api_exec", "form_gen", "analytics", "escalate"}:
        last_user_lower = last_user.lower()
        if any(word in last_user_lower for word in ["document", "file", "pdf", "text", "knowledge"]):
            return "doc_qa"
        elif any(word in last_user_lower for word in ["weather", "search", "api", "get", "fetch"]):
            return "api_exec"
        elif any(word in last_user_lower for word in ["form", "input", "collect", "survey", "field"]):
            return "form_gen"
        elif any(word in last_user_lower for word in ["analyze", "analytics", "statistics", "metrics", "report", "insights", "data"]):
            return "analytics"
        elif any(word in last_user_lower for word in ["human", "agent", "help", "support", "escalate"]):
            return "escalate"
        else:
            return "greeting"
    
    return label


def node_greeting(state: MessagesState):
    llm = get_llm(temperature=0.6)
    sys = (
        "You are a helpful generalist assistant. Be concise and friendly."
    )
    res = llm.invoke([("system", sys), *state["messages"]])
    return {"messages": [res]}


def node_doc_qa(state: MessagesState):
    """Enhanced Document Q&A with chat context memory and multiple document support."""
    tenant_id = CURRENT_TENANT_ID or "default"
    session_id = CURRENT_SESSION.session_id if CURRENT_SESSION else "default"

    # Check if documents are available
    retr = get_retriever_for_tenant(tenant_id)
    documents = document_storage.get_documents_by_tenant(tenant_id)

    if retr is None or not documents:
        content = (
            f"No documents indexed for tenant '{tenant_id}'. "
            f"Please upload documents first using the upload area."
        )
        # Save assistant message to chat history
        save_chat_message_to_history(session_id, tenant_id, "assistant", content, "doc_qa")
        return {"messages": [("assistant", content)]}

    # Find the latest user message
    user_msg = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    # Save user message to chat history
    save_chat_message_to_history(session_id, tenant_id, "user", user_msg, "doc_qa")

    # Get chat history for context
    chat_history = document_storage.get_chat_history(session_id, limit=10)

    # Retrieve relevant documents
    docs = retr(user_msg, k=6)  # Increased for better coverage

    # Group documents by source for better organization
    doc_sources = {}
    for doc in docs:
        source = doc.metadata.get('source', 'Unknown')
        if source not in doc_sources:
            doc_sources[source] = []
        doc_sources[source].append(doc.page_content)

    # Build context with document information
    context_parts = []
    referenced_docs = []

    for source, contents in doc_sources.items():
        # Find document metadata for this source
        doc_info = None
        for doc_meta in documents:
            if doc_meta.file_path.endswith(source) or doc_meta.filename == source:
                doc_info = doc_meta
                referenced_docs.append(doc_meta.document_id)
                break

        doc_name = doc_info.filename if doc_info else source
        context_parts.append(f"[Document: {doc_name}]")
        context_parts.extend(contents)
        context_parts.append("")  # Empty line between documents

    context = "\n".join(context_parts)

    # Build conversation context from recent chat history
    conversation_context = ""
    if len(chat_history) > 1:  # More than just the current message
        recent_messages = chat_history[-6:-1]  # Last 5 messages before current
        conversation_context = "\n".join([
            f"{msg.role.title()}: {msg.content}"
            for msg in recent_messages
        ])

    # Enhanced prompt with conversation context
    prompt_parts = [
        "You are a helpful document Q&A assistant. Answer questions based on the provided documents.",
        "Use the conversation history to maintain context and provide coherent responses.",
        "If the answer is not in the documents, say you don't have enough information.",
        ""
    ]

    if conversation_context:
        prompt_parts.extend([
            "Recent conversation:",
            conversation_context,
            ""
        ])

    prompt_parts.extend([
        "Available documents:",
        context,
        "",
        f"Current question: {user_msg}",
        "",
        "Please provide a comprehensive answer based on the documents above."
    ])

    prompt = "\n".join(prompt_parts)

    # Generate response
    llm = get_llm(temperature=0.1)  # Slightly higher temperature for more natural responses
    res = llm.invoke([("system", "Document QA mode with conversation context."), ("user", prompt)])

    # Save assistant response to chat history
    response_content = getattr(res, "content", str(res))
    save_chat_message_to_history(session_id, tenant_id, "assistant", response_content, "doc_qa", referenced_docs)

    return {"messages": [res]}

def save_chat_message_to_history(session_id: str, tenant_id: str, role: str, content: str,
                                agent_type: Optional[str] = None, document_references: Optional[List[str]] = None):
    """Save a chat message to the persistent chat history."""
    try:
        message = ChatMessage(
            message_id=secrets.token_urlsafe(16),
            session_id=session_id,
            tenant_id=tenant_id,
            role=role,
            content=content,
            timestamp=datetime.now().isoformat(),
            user_id=CURRENT_SESSION.user_id if CURRENT_SESSION else None,
            agent_type=agent_type,
            document_references=document_references or []
        )
        document_storage.save_chat_message(message)
    except Exception as e:
        logger.error(f"Failed to save chat message: {e}")


def node_api_exec(state: MessagesState):
    """Enhanced API execution node with proper tool handling."""
    tenant_id = CURRENT_TENANT_ID or "default"
    tools = get_tenant_tools(tenant_id)
    
    # Create LLM with tools
    llm_with_tools = get_llm(temperature=0).bind_tools(tools)
    
    # Enhanced system prompt for API execution
    system_prompt = (
        "You are an API execution specialist. Your role is to:\n"
        "1. Understand what the user wants to accomplish\n"
        "2. Select and use the appropriate tools to fulfill their request\n"
        "3. Provide clear, helpful responses based on tool results\n"
        "4. Handle errors gracefully and suggest alternatives\n\n"
        f"Available tools: {', '.join([t.name for t in tools])}\n"
        "Always explain what you're doing and why."
    )
    
    # Prepare messages with system prompt
    messages = [("system", system_prompt)] + state["messages"]
    
    # Invoke LLM with tools
    response = llm_with_tools.invoke(messages)
    
    # Check if tools were called
    if hasattr(response, 'tool_calls') and response.tool_calls:
        # Handle tool calls
        tool_node = ToolNode(tools)
        tool_results = tool_node.invoke({"messages": [response]})
        
        # Generate final response with tool results
        final_messages = messages + [response] + tool_results["messages"]
        final_response = get_llm(temperature=0).invoke(final_messages + [
            ("system", "Summarize the results and provide a helpful response to the user.")
        ])
        
        return {"messages": [final_response]}
    
    return {"messages": [response]}


def node_form_gen(state: MessagesState):
    """Professional form generation with PDF/DOC export capabilities."""
    if not has_permission("generate_forms"):
        return {"messages": [("assistant", "Permission denied: form generation not allowed")]}

    llm = get_llm(temperature=0)
    user_msg = ""
    for msg in reversed(state["messages"]):
        # Handle both tuple format ("user", content) and object format
        if isinstance(msg, tuple) and len(msg) >= 2:
            role, content = msg[0], msg[1]
            if role in ["user", "human"]:
                user_msg = content
                break
        elif getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    # Enhanced file format detection
    file_format = "pdf"  # default
    user_msg_lower = user_msg.lower()

    # Check for explicit format specification
    if "(format: docx)" in user_msg_lower or "(format: doc)" in user_msg_lower:
        file_format = "docx"
    elif "(format: pdf)" in user_msg_lower:
        file_format = "pdf"
    # Check for format keywords
    elif any(word in user_msg_lower for word in ["docx", "doc", "word document", "microsoft word"]):
        file_format = "docx"
    elif any(word in user_msg_lower for word in ["pdf", "portable document"]):
        file_format = "pdf"

    # Log the detected format for debugging
    logger.info(f"Full message received: '{user_msg}'")
    logger.info(f"Message lowercase: '{user_msg_lower}'")
    logger.info(f"Contains '(format: docx)': {'(format: docx)' in user_msg_lower}")
    logger.info(f"Contains '(format: pdf)': {'(format: pdf)' in user_msg_lower}")
    logger.info(f"Detected file format: {file_format}")
    
    enhanced_prompt = (
        "You are a professional form generation specialist. Create a comprehensive, structured form based on the user's request.\n"
        "Generate a detailed JSON response with this EXACT structure:\n\n"
        "{\n"
        '  "title": "Professional Form Title",\n'
        '  "description": "Detailed description of the form purpose and instructions",\n'
        '  "company_name": "Company/Organization Name (if applicable)",\n'
        '  "form_type": "contract|survey|registration|feedback|application|contact|other",\n'
        '  "sections": [\n'
        "    {\n"
        '      "title": "Section Title",\n'
        '      "description": "Section description",\n'
        '      "fields": [\n'
        "        {\n"
        '          "name": "field_name",\n'
        '          "label": "Field Label",\n'
        '          "field_type": "text|email|number|date|select|textarea|checkbox|radio|tel",\n'
        '          "required": true|false,\n'
        '          "placeholder": "Placeholder text",\n'
        '          "description": "Field description/help text",\n'
        '          "options": ["option1", "option2"] // Only for select/radio/checkbox\n'
        "        }\n"
        "      ]\n"
        "    }\n"
        "  ],\n"
        '  "footer_text": "Footer text, terms, or additional information"\n'
        "}\n\n"
        "IMPORTANT GUIDELINES:\n"
        "- Create logical sections to organize related fields\n"
        "- Include comprehensive field descriptions and help text\n"
        "- Use appropriate field types for data validation\n"
        "- Add relevant options for select/radio/checkbox fields\n"
        "- Make forms professional and user-friendly\n"
        "- Include proper legal disclaimers for contracts\n"
        "- Add contact information sections where appropriate\n"
        f"\nUser request: {user_msg}"
    )
    
    try:
        # Use proper message format for Gemini
        messages = [
            ("system", "You are a professional form generation specialist. Create comprehensive, structured forms."),
            ("user", enhanced_prompt)
        ]
        res = llm.invoke(messages)
        content = getattr(res, "content", "")

        # Clean up the content to extract JSON
        content = content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]

        # Find JSON boundaries
        start = content.find("{")
        end = content.rfind("}") + 1

        if start != -1 and end > start:
            form_json = content[start:end]

            try:
                form_data = json.loads(form_json)

                # Convert JSON to ProfessionalForm object
                professional_form = _json_to_professional_form(form_data)

                # Generate the file
                try:
                    if file_format == "docx":
                        if not DOCX_AVAILABLE:
                            file_path = FORM_GENERATOR.create_pdf_form(professional_form)
                            file_format = "pdf"  # fallback
                        else:
                            file_path = FORM_GENERATOR.create_docx_form(professional_form)
                    else:
                        file_path = FORM_GENERATOR.create_pdf_form(professional_form)

                    # Create response
                    response_text = (
                        f"✅ **Professional {professional_form.form_type.title()} Form Generated Successfully!**\n\n"
                        f"**Form Details:**\n"
                        f"• Title: {professional_form.title}\n"
                        f"• Type: {professional_form.form_type.title()}\n"
                        f"• Sections: {len(professional_form.sections)}\n"
                        f"• Total Fields: {sum(len(section.fields) for section in professional_form.sections)}\n"
                        f"• Format: {file_format.upper()}\n"
                        f"• File: {file_path}\n"
                        f"• Form ID: {professional_form.form_id}\n\n"
                        f"**Form Structure:**\n"
                    )

                    for i, section in enumerate(professional_form.sections, 1):
                        response_text += f"{i}. **{section.title}** ({len(section.fields)} fields)\n"
                        for field in section.fields[:3]:  # Show first 3 fields
                            required_mark = " *" if field.required else ""
                            response_text += f"   • {field.label}{required_mark} ({field.field_type})\n"
                        if len(section.fields) > 3:
                            response_text += f"   • ... and {len(section.fields) - 3} more fields\n"

                    response_text += (
                        f"\n**File Location:** `{file_path}`\n"
                        f"**Ready for:** Printing, digital distribution, or integration\n\n"
                        f"The form has been professionally formatted with proper headings, sections, "
                        f"field labels, and validation requirements. You can now use this form for "
                        f"your business needs!"
                    )

                    logger.info(f"Generated professional {file_format.upper()} form: {file_path}")
                    return {"messages": [("assistant", response_text)]}

                except Exception as file_error:
                    logger.error(f"File generation error: {file_error}")
                    # Fallback to JSON response
                    form_data["generated_at"] = datetime.now().isoformat()
                    form_data["tenant_id"] = CURRENT_TENANT_ID or "default"
                    form_data["form_id"] = professional_form.form_id

                    fallback_response = (
                        f"⚠️ Form structure generated successfully, but file creation failed.\n"
                        f"Error: {file_error}\n\n"
                        f"**Form JSON Structure:**\n```json\n{json.dumps(form_data, indent=2)}\n```"
                    )
                    return {"messages": [("assistant", fallback_response)]}

            except (json.JSONDecodeError, ValueError) as e:
                logger.error(f"Form parsing error: {e}")
                return {"messages": [("assistant", f"Error parsing form structure: {e}")]}

        # Fallback if no valid JSON found
        return {"messages": [("assistant", "Unable to generate form structure. Please provide more specific requirements.")]}

    except Exception as exc:
        logger.error(f"Form generation error: {exc}")
        return {"messages": [("assistant", f"Error generating form: {exc}")]}


def node_escalate(state: MessagesState):
    """Enhanced escalation workflow with proper handling."""
    try:
        # Get user message for context
        user_msg = ""
        for msg in reversed(state["messages"]):
            if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
                user_msg = getattr(msg, "content", "")
                break
        
        # Create escalation record
        escalation_id = secrets.token_urlsafe(8)
        escalation_data = {
            "escalation_id": escalation_id,
            "tenant_id": CURRENT_TENANT_ID or "default",
            "session_id": CURRENT_SESSION.session_id if CURRENT_SESSION else None,
            "user_message": user_msg,
            "timestamp": datetime.now().isoformat(),
            "status": "pending",
            "conversation_history": [
                {
                    "role": getattr(msg, "type", "unknown"),
                    "content": getattr(msg, "content", str(msg))
                }
                for msg in state["messages"][-5:]  # Last 5 messages for context
            ]
        }
        
        # Log escalation
        logger.info(f"Escalation created: {escalation_id} for tenant {CURRENT_TENANT_ID}")
        
        # In a real system, you would:
        # 1. Store escalation in database
        # 2. Notify human agents
        # 3. Create ticket in support system
        # 4. Send email/slack notification
        
        response = (
            f"I've escalated your request to a human agent.\n\n"
            f"Escalation ID: {escalation_id}\n"
            f"Your request has been logged and a human agent will assist you shortly.\n"
            f"Please keep this escalation ID for reference.\n\n"
            f"In the meantime, you can continue using the chatbot for other queries."
        )
        
        return {"messages": [("assistant", response)]}
        
    except Exception as exc:
        logger.error(f"Escalation error: {exc}")
        return {"messages": [("assistant", "I apologize, but I'm having trouble escalating your request. Please try again or contact support directly.")]}


def node_analytics(state: MessagesState):
    """Analytics agent for data analysis and insights."""
    if not has_permission("use_tools"):
        return {"messages": [("assistant", "Permission denied: analytics not allowed")]}

    llm = get_llm(temperature=0)
    user_msg = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    try:
        # Get system statistics for analysis
        stats = get_system_stats()
        tool_stats = get_tool_stats()

        # Enhanced analytics prompt
        analytics_prompt = (
            "You are an analytics specialist. Analyze the provided system data and user request to provide insights.\n\n"
            f"System Statistics:\n{json.dumps(stats, indent=2)}\n\n"
            f"Tool Usage Statistics:\n{json.dumps(tool_stats, indent=2)}\n\n"
            "Provide detailed analysis including:\n"
            "1. Key metrics and trends\n"
            "2. Usage patterns and insights\n"
            "3. Recommendations for optimization\n"
            "4. Performance indicators\n"
            "5. Actionable insights\n\n"
            f"User request: {user_msg}"
        )

        messages = [
            ("system", "You are a data analytics expert. Provide comprehensive analysis and actionable insights."),
            ("user", analytics_prompt)
        ]

        res = llm.invoke(messages)
        content = getattr(res, "content", "")

        # Enhance response with visual indicators
        enhanced_response = (
            "📊 **Analytics Report**\n\n"
            f"{content}\n\n"
            "---\n"
            f"**Report Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"**Tenant:** {CURRENT_TENANT_ID or 'default'}\n"
            f"**Data Points Analyzed:** {len(stats) + len(tool_stats)}"
        )

        logger.info(f"Generated analytics report for tenant {CURRENT_TENANT_ID}")
        return {"messages": [("assistant", enhanced_response)]}

    except Exception as exc:
        logger.error(f"Analytics error: {exc}")
        return {"messages": [("assistant", f"Error generating analytics: {exc}")]}


# -----------------------------
# Build Enhanced LangGraph
# -----------------------------

def should_continue(state: MessagesState) -> str:
    """Determine if we should continue processing or end."""
    last_message = state["messages"][-1]

    # Check if the last message has tool calls
    if hasattr(last_message, 'tool_calls') and last_message.tool_calls:
        return "tools"

    # Always end after processing (no re-routing)
    return "end"

def create_enhanced_workflow():
    """Create the enhanced multi-agent workflow with memory and tool handling."""
    workflow = StateGraph(MessagesState)

    # Add all agent nodes (no router node, just conditional routing)
    workflow.add_node("greeting", node_greeting)
    workflow.add_node("doc_qa", node_doc_qa)
    workflow.add_node("api_exec", node_api_exec)
    workflow.add_node("form_gen", node_form_gen)
    workflow.add_node("analytics", node_analytics)
    workflow.add_node("escalate", node_escalate)

    # Add tool execution node that dynamically gets tools for current tenant
    def tool_node_func(state: MessagesState):
        tenant_id = CURRENT_TENANT_ID or "default"
        tools = get_tenant_tools(tenant_id)
        tool_node = ToolNode(tools)
        return tool_node.invoke(state)

    workflow.add_node("tools", tool_node_func)

    # Define routing logic - route directly from START
    workflow.add_conditional_edges(
        START,
        node_router,  # Use router as conditional function only
        {
            "greeting": "greeting",
            "doc_qa": "doc_qa",
            "api_exec": "api_exec",
            "form_gen": "form_gen",
            "analytics": "analytics",
            "escalate": "escalate"
        },
    )
    
    # Add conditional edges for tool handling in api_exec
    workflow.add_conditional_edges(
        "api_exec",
        should_continue,
        {"tools": "tools", "end": END}
    )
    
    # Tool results go back to api_exec for final processing
    workflow.add_edge("tools", "api_exec")
    
    # Simple end edges for other agents
    workflow.add_edge("greeting", END)
    workflow.add_edge("doc_qa", END)
    workflow.add_edge("form_gen", END)
    workflow.add_edge("analytics", END)
    workflow.add_edge("escalate", END)
    
    # Add memory for conversation state
    memory = MemorySaver()
    return workflow.compile(checkpointer=memory)

agent = create_enhanced_workflow()


# -----------------------------
# CLI Helpers
# -----------------------------


# -----------------------------
# Admin Dashboard Functions
# -----------------------------

def get_system_stats() -> Dict[str, Any]:
    """Get comprehensive system statistics."""
    stats = {
        "tenants": {
            "total": len(_tenant_registry),
            "active": len([t for t in _tenant_registry.values() if t.is_active]),
            "list": list(_tenant_registry.keys())
        },
        "sessions": {
            "active": len(_active_sessions),
            "total_created": len(_active_sessions)  # Simplified for demo
        },
        "tools": get_tool_stats(),
        "documents": {}
    }
    
    # Get document stats for each tenant
    for tenant_id in _tenant_registry.keys():
        doc_stats = get_document_stats(tenant_id)
        if "error" not in doc_stats:
            stats["documents"][tenant_id] = doc_stats
    
    return stats

def create_admin_dashboard() -> str:
    """Generate admin dashboard HTML."""
    try:
        stats = get_system_stats()

        # Generate components safely
        tenant_rows = _generate_tenant_rows(stats)
        tool_rows = _generate_tool_rows(stats["tools"])

        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Multi-Agent Chatbot Admin Dashboard</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; }}
        .card {{ background: white; padding: 20px; margin: 10px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .stats-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; }}
        .stat-item {{ text-align: center; }}
        .stat-number {{ font-size: 2em; font-weight: bold; color: #2196F3; }}
        .stat-label {{ color: #666; margin-top: 5px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
        th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background-color: #f8f9fa; }}
        .status-active {{ color: #4CAF50; font-weight: bold; }}
        .status-inactive {{ color: #f44336; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Multi-Agent Chatbot Admin Dashboard</h1>

        <div class="stats-grid">
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["tenants"]["total"]}</div>
                    <div class="stat-label">Total Tenants</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["tenants"]["active"]}</div>
                    <div class="stat-label">Active Tenants</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["sessions"]["active"]}</div>
                    <div class="stat-label">Active Sessions</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{len(stats["tools"])}</div>
                    <div class="stat-label">Total Tools</div>
                </div>
            </div>
        </div>

        <div class="card">
            <h2>Tenant Overview</h2>
            <table>
                <thead>
                    <tr>
                        <th>Tenant ID</th>
                        <th>Name</th>
                        <th>Status</th>
                        <th>Documents</th>
                        <th>Permissions</th>
                    </tr>
                </thead>
                <tbody>
                    {tenant_rows}
                </tbody>
            </table>
        </div>

        <div class="card">
            <h2>Tool Usage Statistics</h2>
            <table>
                <thead>
                    <tr>
                        <th>Tool Name</th>
                        <th>Total Calls</th>
                        <th>Errors</th>
                        <th>Success Rate</th>
                    </tr>
                </thead>
                <tbody>
                    {tool_rows}
                </tbody>
            </table>
        </div>
    </div>
</body>
</html>"""

        return html
    except Exception as e:
        logger.error(f"Dashboard generation error: {e}")
        return f"<html><body><h1>Dashboard Error</h1><p>Error: {e}</p></body></html>"

def _generate_tenant_rows(stats: Dict) -> str:
    """Generate HTML rows for tenant table."""
    rows = []
    for tenant_id, config in _tenant_registry.items():
        doc_count = stats["documents"].get(tenant_id, {}).get("total_chunks", 0)
        status_class = "status-active" if config.is_active else "status-inactive"
        status_text = "Active" if config.is_active else "Inactive"
        
        row = f"""
        <tr>
            <td>{tenant_id}</td>
            <td>{config.name}</td>
            <td class="{status_class}">{status_text}</td>
            <td>{doc_count} chunks</td>
            <td>{', '.join(config.permissions[:3])}{'...' if len(config.permissions) > 3 else ''}</td>
        </tr>
        """
        rows.append(row)
    
    return "".join(rows)

def _generate_tool_rows(tool_stats: Dict) -> str:
    """Generate HTML rows for tool statistics table."""
    rows = []
    for tool_name, data in tool_stats.items():
        calls = data['call_count']
        errors = data['error_count']
        success_rate = ((calls - errors) / calls * 100) if calls > 0 else 0
        
        row = f"""
        <tr>
            <td>{tool_name}</td>
            <td>{calls}</td>
            <td>{errors}</td>
            <td>{success_rate:.1f}%</td>
        </tr>
        """
        rows.append(row)
    
    return "".join(rows)

# -----------------------------
# Enhanced CLI Commands
# -----------------------------

def handle_command(line: str) -> Optional[str]:
    global CURRENT_TENANT_ID, CURRENT_SESSION
    
    if line.startswith("/tenant "):
        tenant_id = line.split(" ", 1)[1].strip() or None
        if tenant_id and authenticate_tenant(tenant_id):
            CURRENT_TENANT_ID = tenant_id
            # Create session for tenant
            try:
                CURRENT_SESSION = create_session(tenant_id)
                return f"Active tenant set to: {CURRENT_TENANT_ID} (Session: {CURRENT_SESSION.session_id[:8]}...)"
            except ValueError as e:
                return f"Error: {e}"
        else:
            return f"Invalid or inactive tenant: {tenant_id}"
    
    if line.startswith("/who"):
        session_info = f" (Session: {CURRENT_SESSION.session_id[:8]}...)" if CURRENT_SESSION else ""
        return f"Active tenant: {CURRENT_TENANT_ID}{session_info}"
    
    if line.startswith("/create-tenant "):
        parts = line.split(" ", 2)
        if len(parts) < 3:
            return "Usage: /create-tenant TENANT_ID TENANT_NAME"
        tenant_id, name = parts[1], parts[2]
        try:
            config = create_tenant(tenant_id, name)
            return f"Created tenant '{tenant_id}' ({name})"
        except ValueError as e:
            return f"Error: {e}"
    
    if line.startswith("/ingest "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("read_documents"):
            return "Permission denied: document ingestion not allowed"
        path = line.split(" ", 1)[1].strip().strip('"')
        return ingest_documents_from_dir(CURRENT_TENANT_ID, path)
    
    if line.startswith("/tool.httpget "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("use_tools"):
            return "Permission denied: tool registration not allowed"
        # Usage: /tool.httpget NAME BASE_URL_ENV [API_KEY_ENV]
        parts = line.split()
        if len(parts) < 3:
            return "Usage: /tool.httpget NAME BASE_URL_ENV [API_KEY_ENV]"
        name = parts[1]
        base_env = parts[2]
        api_env = parts[3] if len(parts) > 3 else None
        t = make_http_get_tool(name=name, description=f"HTTP GET tool for {name}", base_url_env=base_env, api_key_env=api_env)
        register_dynamic_tool(CURRENT_TENANT_ID, t)
        return f"Registered tool '{name}' for tenant {CURRENT_TENANT_ID}."
    
    if line.startswith("/tool.httppost "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("use_tools"):
            return "Permission denied: tool registration not allowed"
        parts = line.split()
        if len(parts) < 3:
            return "Usage: /tool.httppost NAME BASE_URL_ENV [API_KEY_ENV]"
        name = parts[1]
        base_env = parts[2]
        api_env = parts[3] if len(parts) > 3 else None
        t = make_http_post_tool(name=name, description=f"HTTP POST tool for {name}", base_url_env=base_env, api_key_env=api_env)
        register_dynamic_tool(CURRENT_TENANT_ID, t)
        return f"Registered POST tool '{name}' for tenant {CURRENT_TENANT_ID}."
    
    if line.startswith("/tools"):
        names = [t.name for t in get_tenant_tools(CURRENT_TENANT_ID)]
        return "Available tools: " + ", ".join(names)
    
    if line.startswith("/stats"):
        if not has_permission("admin"):
            return get_document_stats_tool()  # Limited stats for non-admin
        stats = get_system_stats()
        result = f"System Statistics:\n"
        result += f"- Tenants: {stats['tenants']['total']} total, {stats['tenants']['active']} active\n"
        result += f"- Active Sessions: {stats['sessions']['active']}\n"
        result += f"- Tools: {len(stats['tools'])} registered\n"
        return result
    
    if line.startswith("/dashboard"):
        if not has_permission("admin"):
            return "Permission denied: admin access required"
        try:
            html = create_admin_dashboard()
            dashboard_file = "admin_dashboard.html"
            with open(dashboard_file, "w", encoding="utf-8") as f:
                f.write(html)
            return f"Admin dashboard saved to {dashboard_file}. Open in browser to view."
        except Exception as e:
            return f"Error creating dashboard: {e}"
    
    if line.startswith("/permissions"):
        if not CURRENT_SESSION:
            return "No active session"
        return f"Your permissions: {', '.join(CURRENT_SESSION.permissions)}"
    
    if line.startswith("/help"):
        return (
            "Available Commands:\n"
            "  /tenant TENANT_ID                    Set active tenant\n"
            "  /create-tenant ID NAME               Create new tenant (admin)\n"
            "  /who                                 Show active tenant and session\n"
            "  /permissions                         Show your permissions\n"
            "  /ingest PATH                         Ingest documents from directory\n"
            "  /tool.httpget NAME BASE_URL_ENV [KEY_ENV]   Register HTTP GET tool\n"
            "  /tool.httppost NAME BASE_URL_ENV [KEY_ENV]  Register HTTP POST tool\n"
            "  /tools                               List available tools\n"
            "  /stats                               Show system statistics\n"
            "  /dashboard                           Generate admin dashboard (admin)\n"
            "  /help                                Show this help"
        )
    
    return None


def chat_once(user_input: str, thread_id: str = "default") -> str:
    """Chat with the agent using proper thread configuration."""
    config = {"configurable": {"thread_id": thread_id}}
    final_state = agent.invoke({"messages": [("user", user_input)]}, config=config)
    last = final_state["messages"][-1]
    return getattr(last, "content", str(last))

def chat_with_agent(user_input: str, tenant_id: str = "default") -> str:
    """Simplified chat function for testing and demos."""
    global CURRENT_TENANT_ID, CURRENT_SESSION
    original_tenant = CURRENT_TENANT_ID
    original_session = CURRENT_SESSION

    try:
        # Set tenant context
        CURRENT_TENANT_ID = tenant_id

        # Create or get session for this tenant
        if tenant_id in _tenant_registry:
            session = create_session(tenant_id)
            CURRENT_SESSION = session

        thread_id = f"chat_{tenant_id}_{hash(user_input) % 1000}"
        result = chat_once(user_input, thread_id)
        return result
    finally:
        CURRENT_TENANT_ID = original_tenant
        CURRENT_SESSION = original_session


if __name__ == "__main__":
    if not os.environ.get("GOOGLE_API_KEY"):
        print("Warning: GOOGLE_API_KEY is not set. Set it to enable the LLM and embeddings.")
    print("Multi-Agent Chatbot (LangGraph + LangChain + FAISS RAG). Type '/help' for commands, 'exit' to quit.\n")
    # Choose tenant at start
    try:
        CURRENT_TENANT_ID = input("Tenant ID (default): ").strip() or "default"
    except (EOFError, KeyboardInterrupt):
        CURRENT_TENANT_ID = "default"
    print(f"Active tenant: {CURRENT_TENANT_ID}\n")
    while True:
        try:
            user = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nBye!")
            break
        if not user:
            continue
        if user.lower() in {"exit", "quit"}:
            print("Bye!")
            break
        if user.startswith("/"):
            out = handle_command(user)
            print(f"Bot: {out}\n")
            continue
        # Generate unique thread ID for conversation
        thread_id = f"session_{CURRENT_TENANT_ID}_{hash(user) % 10000}"
        reply = chat_once(user, thread_id)
        print(f"Bot: {reply}\n")